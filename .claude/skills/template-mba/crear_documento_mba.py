"""
TEMPLATE MBA - FEN UCHILE v5.1
================================

Script optimizado para crear documentos académicos MBA con formato aprobado por FEN UChile.

CARACTERÍSTICAS:
- Logo FEN UChile en encabezado (ancho completo optimizado)
- Tabla de contenidos automática con título
- Numeración automática decimal (1., 2., 2.1, 2.1.1)
- Formato de portada correcto (cursiva y color azul)
- Tablas con formato MBA estándar
- Pie de página con numeración
- Procesamiento correcto de listas con espaciado adecuado

AUTOR: Luis Rivera González + Claude AI
VERSIÓN: 5.1
FECHA: 25 de Octubre 2025
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


def add_toc(doc):
    """
    Inserta un campo TOC (Table of Contents) con título.
    El TOC se actualiza automáticamente en Word con Ctrl+A + F9
    """
    # Agregar título "TABLA DE CONTENIDO"
    titulo_toc = doc.add_paragraph('TABLA DE CONTENIDO')
    titulo_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_toc.runs[0].font.name = 'Arial'
    titulo_toc.runs[0].font.size = Pt(14)
    titulo_toc.runs[0].font.bold = True
    titulo_toc.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    # Espacio después del título
    doc.add_paragraph()
    
    # Insertar campo TOC
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    # Estructura OOXML del campo TOC
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    t = OxmlElement('w:t')
    t.text = 'Presione Ctrl+A y luego F9 para actualizar la tabla de contenido'
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(t)
    run._r.append(fldChar3)
    
    doc.add_paragraph()
    return paragraph


def add_page_number(paragraph):
    """Agrega numeración de página al pie de página"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def create_multilevel_numbering(doc):
    """
    Crea una numeración multinivel DECIMAL (sin bullets) en el documento.
    
    Configuración:
    - Nivel 0 (Heading 2): 1., 2., 3., ...
    - Nivel 1 (Heading 3): 1.1, 1.2, 2.1, ...
    - Nivel 2 (Heading 4): 1.1.1, 1.1.2, ...
    
    Returns:
        str: numId de la numeración creada
    """
    # Obtener numbering part
    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        raise ValueError("El documento no tiene numbering part")
    
    # Crear abstractNum (definición de numeración)
    abstractNum = OxmlElement('w:abstractNum')
    abstractNum.set(qn('w:abstractNumId'), '2')
    
    # Configuración de niveles
    niveles = [
        {  # Nivel 0: Heading 2 → 1., 2., 3., ...
            'ilvl': '0',
            'numFmt': 'decimal',
            'lvlText': '%1.',
            'lvlJc': 'left',
            'pPr_ind_left': '720',
            'pPr_ind_hanging': '360'
        },
        {  # Nivel 1: Heading 3 → 1.1, 1.2, ...
            'ilvl': '1',
            'numFmt': 'decimal',
            'lvlText': '%1.%2',
            'lvlJc': 'left',
            'pPr_ind_left': '1080',
            'pPr_ind_hanging': '360'
        },
        {  # Nivel 2: Heading 4 → 1.1.1, 1.1.2, ...
            'ilvl': '2',
            'numFmt': 'decimal',
            'lvlText': '%1.%2.%3',
            'lvlJc': 'left',
            'pPr_ind_left': '1440',
            'pPr_ind_hanging': '360'
        }
    ]
    
    # Crear cada nivel
    for nivel_config in niveles:
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), nivel_config['ilvl'])
        
        # Número de inicio
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)
        
        # Formato de número
        numFmt = OxmlElement('w:numFmt')
        numFmt.set(qn('w:val'), nivel_config['numFmt'])
        lvl.append(numFmt)
        
        # Texto del nivel
        lvlText = OxmlElement('w:lvlText')
        lvlText.set(qn('w:val'), nivel_config['lvlText'])
        lvl.append(lvlText)
        
        # Alineación
        lvlJc = OxmlElement('w:lvlJc')
        lvlJc.set(qn('w:val'), nivel_config['lvlJc'])
        lvl.append(lvlJc)
        
        # Propiedades de párrafo (indentación)
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), nivel_config['pPr_ind_left'])
        ind.set(qn('w:hanging'), nivel_config['pPr_ind_hanging'])
        pPr.append(ind)
        lvl.append(pPr)
        
        # Propiedades de fuente
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')
        rPr.append(rFonts)
        lvl.append(rPr)
        
        abstractNum.append(lvl)
    
    # Agregar al numbering.xml
    numbering_part._element.append(abstractNum)
    
    # Crear num que referencia este abstractNum
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), '2')
    abstractNumId = OxmlElement('w:abstractNumId')
    abstractNumId.set(qn('w:val'), '2')
    num.append(abstractNumId)
    numbering_part._element.append(num)
    
    return '2'


def add_numbering_to_paragraph(paragraph, numId, ilvl):
    """
    Aplica numeración a un párrafo específico.
    
    Args:
        paragraph: Párrafo de docx
        numId: ID de la numeración (normalmente '2')
        ilvl: Nivel de indentación (0, 1, 2, etc.)
    """
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    
    # Nivel de numeración
    ilvl_elem = OxmlElement('w:ilvl')
    ilvl_elem.set(qn('w:val'), str(ilvl))
    numPr.append(ilvl_elem)
    
    # ID de numeración
    numId_elem = OxmlElement('w:numId')
    numId_elem.set(qn('w:val'), str(numId))
    numPr.append(numId_elem)
    
    pPr.append(numPr)


def configurar_estilos_heading(doc):
    """
    Configura los estilos de Heading con numeración automática MBA
    """
    styles = doc.styles
    
    # Heading 2: 1., 2., 3., ... (capítulos principales)
    heading2 = styles['Heading 2']
    heading2.font.name = 'Arial'
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0, 0, 0)
    
    # Heading 3: 1.1, 1.2, ... (subcapítulos)
    heading3 = styles['Heading 3']
    heading3.font.name = 'Arial'
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3.font.color.rgb = RGBColor(0, 0, 0)
    
    # Heading 4: 1.1.1, 1.1.2, ... (sub-subcapítulos)
    heading4 = styles['Heading 4']
    heading4.font.name = 'Arial'
    heading4.font.size = Pt(11)
    heading4.font.bold = True
    heading4.font.color.rgb = RGBColor(0, 0, 0)


def crear_tabla_mba(doc, encabezados, datos, ancho_columnas=None):
    """
    Crea una tabla con formato MBA estándar:
    - Encabezado con fondo gris (D9D9D9)
    - Texto centrado en encabezado, negrita
    - Texto alineado a la izquierda en datos
    - Bordes negros en todas las celdas
    - Fuente Arial
    
    Args:
        doc: Documento docx
        encabezados: Lista de strings con nombres de columnas
        datos: Lista de listas con datos de cada fila
        ancho_columnas: Lista opcional con anchos en pulgadas
    
    Returns:
        Table object
        
    Ejemplo:
        encabezados = ['Criterio', 'Ponderación', 'Resultado']
        datos = [
            ['VAN', 'Alta', '$7,143K'],
            ['TIR', 'Alta', '15.58 pp']
        ]
        crear_tabla_mba(doc, encabezados, datos)
    """
    # Crear tabla
    num_cols = len(encabezados)
    num_rows = len(datos) + 1
    
    tabla = doc.add_table(rows=num_rows, cols=num_cols)
    
    # Función para configurar bordes
    def set_cell_border(cell):
        """Configura bordes negros en una celda"""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        
        tcPr.append(tcBorders)
    
    # Configurar anchos si se especifican
    if ancho_columnas:
        for i, ancho in enumerate(ancho_columnas):
            for cell in tabla.columns[i].cells:
                cell.width = Inches(ancho)
    
    # ENCABEZADO
    header_cells = tabla.rows[0].cells
    for i, texto in enumerate(encabezados):
        cell = header_cells[i]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(texto)
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Fondo gris
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D9D9D9')
        cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Bordes
        set_cell_border(cell)
    
    # DATOS
    for i, fila_datos in enumerate(datos, start=1):
        row_cells = tabla.rows[i].cells
        for j, valor in enumerate(fila_datos):
            cell = row_cells[j]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(str(valor))
            run.font.name = 'Arial'
            run.font.size = Pt(10.5)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Bordes
            set_cell_border(cell)
    
    # Márgenes de celda
    for row in tabla.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for margin_name in ['top', 'left', 'bottom', 'right']:
                node = OxmlElement(f'w:{margin_name}')
                node.set(qn('w:w'), '100')
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)
    
    return tabla


def procesar_contenido_con_listas(doc, contenido_texto):
    """
    Procesa contenido que puede contener listas con viñetas (•)
    sin agregar espaciado excesivo.
    
    Args:
        doc: Documento docx
        contenido_texto: Texto que puede contener listas con formato:
            "Texto normal
            
            • Item 1.
            • Item 2.
            
            Más texto normal"
    """
    # Dividir por doble salto de línea (párrafos)
    parrafos = contenido_texto.split('\n\n')
    
    for parrafo in parrafos:
        parrafo = parrafo.strip()
        if not parrafo:
            continue
            
        lineas = parrafo.split('\n')
        
        # Verificar si es una lista (todas las líneas con •)
        es_lista = all(l.strip().startswith('•') for l in lineas if l.strip())
        
        if es_lista:
            # Procesar como lista
            for linea in lineas:
                linea = linea.strip()
                if linea:
                    p = doc.add_paragraph(linea)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    # Formato de fuente
                    for run in p.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10.5)
        else:
            # Procesar como párrafo normal (sin espacio adicional)
            p = doc.add_paragraph(parrafo)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10.5)


def crear_documento_mba(template_path, output_path, contenido):
    """
    Crea un documento MBA completo a partir del template.
    
    Args:
        template_path: Ruta al TEMPLATE_MBA_OPTIMIZADO.docx
        output_path: Ruta donde guardar el documento final
        contenido: Dict con la siguiente estructura:
            {
                'curso': str,           # Nombre del curso
                'ejercicio': str,       # Título del ejercicio/caso
                'fecha': str,           # Fecha (DD/MM/AAAA)
                'capitulos': [          # Lista de capítulos
                    {
                        'titulo': str,
                        'nivel': int,   # 2, 3, o 4 (Heading level)
                        'contenido': str,  # Puede contener listas con •
                        'tablas': [     # Opcional
                            {
                                'encabezados': [...],
                                'datos': [[...]],
                                'anchos': [...] (opcional)
                            }
                        ]
                    }
                ]
            }
    
    Returns:
        None (guarda el archivo)
    """
    # Cargar template (el logo ya está optimizado en el template)
    doc = Document(template_path)
    configurar_estilos_heading(doc)
    
    # Eliminar contenido después de integrantes (párrafo 21)
    for i in range(len(doc.paragraphs) - 1, 21, -1):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)
    
    # Personalizar portada con FORMATO CORRECTO
    # Párrafo 8: CURSO (mayúsculas, negrita, Arial 12pt)
    p_curso = doc.paragraphs[8]
    p_curso.clear()
    run_curso = p_curso.add_run(contenido.get('curso', 'NOMBRE DEL CURSO').upper())
    run_curso.font.name = 'Arial'
    run_curso.font.size = Pt(12)
    run_curso.font.bold = True
    run_curso.font.color.rgb = RGBColor(0, 0, 0)
    p_curso.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Párrafo 10: EJERCICIO (cursiva, azul #2E5894, Arial 11pt)
    p_ejercicio = doc.paragraphs[10]
    p_ejercicio.clear()
    run_ejercicio = p_ejercicio.add_run(f"Ejercicio {contenido.get('ejercicio', '[Título]')}")
    run_ejercicio.font.name = 'Arial'
    run_ejercicio.font.size = Pt(11)
    run_ejercicio.font.italic = True
    run_ejercicio.font.color.rgb = RGBColor(46, 88, 148)  # #2E5894
    p_ejercicio.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Párrafo 11: FECHA (azul #2E5894, Arial 11pt)
    p_fecha = doc.paragraphs[11]
    p_fecha.clear()
    run_fecha = p_fecha.add_run(f"FECHA: {contenido.get('fecha', 'DD/MM/AAAA')}")
    run_fecha.font.name = 'Arial'
    run_fecha.font.size = Pt(11)
    run_fecha.font.color.rgb = RGBColor(46, 88, 148)  # #2E5894
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Salto de página + TOC
    doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
    add_toc(doc)
    doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
    
    # Crear sistema de numeración multinivel
    numId = create_multilevel_numbering(doc)
    
    # Agregar capítulos con numeración automática
    for cap in contenido.get('capitulos', []):
        # Título del capítulo
        heading = doc.add_heading(cap['titulo'], level=cap.get('nivel', 2))
        
        # Aplicar numeración según el nivel
        nivel = cap.get('nivel', 2)
        if nivel == 2:
            add_numbering_to_paragraph(heading, numId, 0)  # Nivel 0: 1., 2., 3.
        elif nivel == 3:
            add_numbering_to_paragraph(heading, numId, 1)  # Nivel 1: 1.1, 1.2
        elif nivel == 4:
            add_numbering_to_paragraph(heading, numId, 2)  # Nivel 2: 1.1.1, 1.1.2
        
        # Contenido del capítulo (con procesamiento de listas)
        if cap.get('contenido'):
            procesar_contenido_con_listas(doc, cap['contenido'])
        
        # Tablas (si existen)
        for tabla_data in cap.get('tablas', []):
            crear_tabla_mba(
                doc,
                tabla_data['encabezados'],
                tabla_data['datos'],
                tabla_data.get('anchos')
            )
            doc.add_paragraph()
    
    # Configurar pie de página con numeración
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer_para)
    
    # Guardar
    doc.save(output_path)
    print(f"✅ Documento guardado: {output_path}")
    print("\n📌 INSTRUCCIONES:")
    print("1. Abrir el documento en Microsoft Word")
    print("2. Presionar Ctrl+A (seleccionar todo)")
    print("3. Presionar F9 (actualizar campos)")
    print("\n✓ La tabla de contenidos se generará automáticamente")


# EJEMPLO DE USO
if __name__ == "__main__":
    
    # Definir el contenido del documento
    mi_contenido = {
        'curso': 'FINANZAS CORPORATIVAS',
        'ejercicio': 'Análisis de Inversión en Proyecto de Expansión',
        'fecha': '25/10/2025',
        'capitulos': [
            {
                'titulo': 'RESUMEN EJECUTIVO',
                'nivel': 2,
                'contenido': 'El presente trabajo analiza la viabilidad financiera de un proyecto de expansión.'
            },
            {
                'titulo': 'ANÁLISIS FINANCIERO',
                'nivel': 2,
                'contenido': 'Se evalúan los principales indicadores de rentabilidad del proyecto.'
            },
            {
                'titulo': 'Estructura de Inversión',
                'nivel': 3,
                'contenido': 'La inversión se distribuye en los siguientes rubros:',
                'tablas': [
                    {
                        'encabezados': ['Rubro', 'Monto (MUSD)', 'Porcentaje'],
                        'datos': [
                            ['Terrenos', '5.2', '8.7%'],
                            ['Maquinaria', '28.5', '47.5%'],
                            ['Construcción', '18.3', '30.5%'],
                            ['Capital trabajo', '8.0', '13.3%']
                        ],
                        'anchos': [3.5, 2.0, 1.5]
                    }
                ]
            },
            {
                'titulo': 'Indicadores de Rentabilidad',
                'nivel': 3,
                'contenido': '''Los principales indicadores del proyecto son:

• VAN @ 12%: USD 125.5 M - VAN > 0 ✓
• TIR: 18.7% - TIR > WACC ✓
• Payback: 4.2 años - Dentro horizonte ✓

Todos los indicadores muestran viabilidad del proyecto.''',
            },
            {
                'titulo': 'CONCLUSIONES',
                'nivel': 2,
                'contenido': 'El proyecto presenta indicadores favorables que justifican su implementación.'
            }
        ]
    }
    
    # Crear el documento
    crear_documento_mba(
        template_path='TEMPLATE_MBA_OPTIMIZADO.docx',
        output_path='Mi_Documento_MBA.docx',
        contenido=mi_contenido
    )
